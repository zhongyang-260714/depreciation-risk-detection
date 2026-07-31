# -*- coding: utf-8 -*-
"""
将挑战杯竞赛报告 Markdown 转换为「标准论文版」Word 文档。
特性：
- A4 / 页边距 2.54/3.18cm，小四宋体正文，1.5 倍行距，首行缩进 2 字符，两端对齐
- 一级标题(章)黑体三号居中+分页；二级黑体四号；三级黑体小四
- 自动生成目录域(打开后更新域即出页码)
- 表格：表头底纹+边框，按内容比例分配列宽适配页面
- 代码块/引用块：底纹+左边框
- 图片：嵌入并加图注
- 页脚居中页码
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\科创企业资产折旧算法\竞赛报告_完整稿_v2.1.md"
IMG_DIR = r"D:\科创企业资产折旧算法\报告图表"
OUT = r"D:\科创企业资产折旧算法\竞赛报告_完整稿_v2.1_论文版.docx"

CN = "宋体"
HEI = "黑体"
MONO = "Consolas"

USABLE_W = Cm(14.6)  # 21 - 3.18 - 3.18

# ---------- 字体/样式辅助 ----------
def set_run_font(run, cn=CN, ascii_font="Times New Roman", size=None, bold=None, color=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), ascii_font)
    rf.set(qn('w:hAnsi'), ascii_font)
    rf.set(qn('w:eastAsia'), cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def set_para_shading(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    pPr.append(shd)

def set_left_border(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '0'); left.set(qn('w:color'), hexcolor)
    pbdr.append(left); pPr.append(pbdr)

def set_cell_shading(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

# ---------- 行内加粗/代码解析 ----------
def add_runs(par, text, base_size=12, cn=CN, ascii_font="Times New Roman", bold_default=None):
    text = text.replace('****', '**')  # 修正 **** 笔误
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for tok in parts:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
            set_run_font(r, cn, ascii_font, base_size)
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1])
            set_run_font(r, MONO, MONO, max(8, base_size - 1))
        else:
            r = par.add_run(tok)
            set_run_font(r, cn, ascii_font, base_size, bold=bold_default)

# ---------- 文档级设置 ----------
def configure_styles(doc):
    normal = doc.styles['Normal']
    normal.font.size = Pt(12)
    normal.font.name = 'Times New Roman'
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
    rf.set(qn('w:eastAsia'), CN)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

def add_page_number_footer(doc):
    sec = doc.sections[0]
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_run_font(run, CN, 'Times New Roman', 10.5)

def add_toc(doc):
    # 目录标题
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目  录'); set_run_font(r, HEI, 'Times New Roman', 16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    # 域
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fldB = OxmlElement('w:fldChar'); fldB.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u \\p'
    fldS = OxmlElement('w:fldChar'); fldS.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '（打开后请右键此处“更新域”，或按 F9 刷新目录与页码）'
    fldE = OxmlElement('w:fldChar'); fldE.set(qn('w:fldCharType'), 'end')
    run._r.append(fldB); run._r.append(instr); run._r.append(fldS); run._r.append(t); run._r.append(fldE)
    set_run_font(run, CN, 'Times New Roman', 10.5, color=RGBColor(0x80,0x80,0x80))
    # 打开时自动更新域
    settings = doc.settings.element
    up = OxmlElement('w:updateFields'); up.set(qn('w:val'), 'true')
    settings.append(up)

# ---------- 各种块 ----------
def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.style = doc.styles['Title']; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size, cn, bold = 22, HEI, True
    elif level == 1:
        p.style = doc.styles['Heading 1']; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        size, cn, bold = 16, HEI, True
    elif level == 2:
        p.style = doc.styles['Heading 2']; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, cn, bold = 14, HEI, True
    elif level == 3:
        p.style = doc.styles['Heading 3']; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, cn, bold = 12, HEI, True
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size, cn, bold = 12, HEI, True
    r = p.add_run(text); set_run_font(r, cn, 'Times New Roman', size, bold=bold)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    return p

def add_body(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, caption=False):
    p = doc.add_paragraph()
    if caption:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    add_runs(p, text, base_size=size)
    return p

def add_callout(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    set_para_shading(p, 'F2F2F2'); set_left_border(p, '2E75B6')
    first = True
    for line in lines:
        if not first:
            p.add_run().add_break()
        first = False
        add_runs(p, line, base_size=10.5)
    return p

def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    set_para_shading(p, 'F5F5F5')
    first = True
    for line in lines:
        if not first:
            p.add_run().add_break()
        first = False
        r = p.add_run(line)
        set_run_font(r, MONO, MONO, 9)
    return p

def add_list(doc, items):
    for (indent, kind, content) in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74 * (indent + 1))
        p.paragraph_format.line_spacing = 1.5
        p.style = doc.styles['List Bullet'] if kind == 'bullet' else doc.styles['List Number']
        add_runs(p, content, base_size=12)
    return

def parse_table(rows):
    data = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue  # 分隔行
        data.append(cells)
    return data

def add_table(doc, data):
    if not data:
        return
    ncol = max(len(r) for r in data)
    for r in data:
        while len(r) < ncol:
            r.append('')
    tbl = doc.add_table(rows=len(data), cols=ncol)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    # 列宽按内容长度比例分配
    colchars = [max(len(r[c]) for r in data) for c in range(ncol)]
    tot = sum(colchars) or 1
    colw = [USABLE_W * (colchars[c] / tot) for c in range(ncol)]
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = colw[ci]
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ri == 0 or len(val) <= 12) else WD_ALIGN_PARAGRAPH.LEFT
            run = cp.add_run(val)
            if ri == 0:
                set_run_font(run, HEI, 'Times New Roman', 10.5, bold=True)
                set_cell_shading(cell, 'DCE6F1')
            else:
                set_run_font(run, CN, 'Times New Roman', 10.5)
    return

def add_image(doc, alt, relpath):
    path = relpath
    if not os.path.isabs(path):
        cand = os.path.join(IMG_DIR, os.path.basename(relpath))
        if os.path.exists(cand):
            path = cand
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        try:
            p.add_run().add_picture(path, width=Cm(14.5))
        except Exception as e:
            r = p.add_run('【图片缺失：%s】' % os.path.basename(path)); set_run_font(r, CN, 'Times New Roman', 10.5)
    else:
        r = p.add_run('【图片缺失：%s】' % relpath); set_run_font(r, CN, 'Times New Roman', 10.5)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(alt); set_run_font(r, CN, 'Times New Roman', 10.5)
    cap.paragraph_format.space_after = Pt(8)
    return

# ---------- 主流程 ----------
def main():
    doc = Document()
    configure_styles(doc)
    add_page_number_footer(doc)

    lines = open(SRC, encoding='utf-8').read().split('\n')

    phase = 'start'            # start -> subtitle -> meta -> body
    toc_pending = False
    toc_emitted = False
    title_done = False
    subtitle_done = False

    bq_buf = []
    list_buf = []
    table_buf = []
    in_code = False
    code_buf = []

    def flush_bq():
        if bq_buf:
            add_callout(doc, [l for l in bq_buf]); bq_buf.clear()
    def flush_list():
        if list_buf:
            add_list(doc, list_buf); list_buf.clear()
    def flush_table():
        if table_buf:
            add_table(doc, parse_table(table_buf)); table_buf.clear()

    def handle_heading(text, level):
        nonlocal phase, toc_pending, toc_emitted, title_done, subtitle_done
        if level == 1 and not title_done:
            # 文档标题(封面)
            add_heading(doc, text, 0)
            title_done = True; phase = 'after_title'; return
        if level == 2 and text.startswith('——') and phase in ('after_title',):
            # 副标题
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text); set_run_font(r, HEI, 'Times New Roman', 14, bold=True)
            p.paragraph_format.space_after = Pt(10)
            subtitle_done = True; phase = 'after_subtitle'; return
        if level == 2 and text.strip() == '目录':
            toc_pending = True; return  # 稍后(摘要之后)生成
        if level == 1 and text.strip() == '摘要':
            add_heading(doc, text, 1)
            phase = 'after_abstract'; return
        # 其他：第一章/参考文献/附录 等一级；其它二级/三级
        if level == 1:
            if toc_pending and not toc_emitted:
                # 在正文首章之前插入目录页
                pb = doc.add_paragraph(); pb.paragraph_format.page_break_before = True
                add_toc(doc); toc_emitted = True
            add_heading(doc, text, 1)
        elif level == 2:
            add_heading(doc, text, 2)
        elif level == 3:
            add_heading(doc, text, 3)
        else:
            add_heading(doc, text, 4)

    for raw in lines:
        line = raw.rstrip('\n')
        stripped = line.strip()
        if stripped == '':
            continue
        if re.match(r'^-{3,}$', stripped):  # 水平分隔线：论文中不渲染
            continue
        if stripped.startswith('```'):
            flush_bq(); flush_list(); flush_table()
            if in_code:
                add_code(doc, code_buf); in_code = False; code_buf = []
            else:
                in_code = True; code_buf = []
            continue
        if in_code:
            code_buf.append(line); continue

        # 标题
        hm = re.match(r'^(#{1,6})\s+(.*)$', line)
        if hm:
            flush_bq(); flush_list(); flush_table()
            handle_heading(hm.group(2).strip(), len(hm.group(1)))
            continue

        # 引用块
        if stripped.startswith('>'):
            flush_list(); flush_table()
            bq_buf.append(re.sub(r'^>\s?', '', stripped))
            continue
        else:
            flush_bq()

        # 表格行
        if stripped.startswith('|'):
            flush_list()
            table_buf.append(stripped); continue
        else:
            flush_table()

        # 图片
        im = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', stripped)
        if im:
            flush_list()
            add_image(doc, im.group(1).strip(), im.group(2).strip())
            continue

        # 列表项
        bm = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        nm = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if bm or nm:
            flush_table()
            if bm:
                indent = len(bm.group(1)) // 2
                list_buf.append((indent, 'bullet', bm.group(2)))
            else:
                indent = len(nm.group(1)) // 2
                list_buf.append((indent, 'number', nm.group(3)))
            continue
        else:
            flush_list()

        # 普通段落（含封面元信息 / 表图注 / 正文）
        if phase in ('after_subtitle', 'meta'):
            # 封面竞赛元信息：居中
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, stripped, base_size=12)
            phase = 'meta'
            continue
        # 表/图注
        if stripped.startswith('**表') or stripped.startswith('**图'):
            add_body(doc, stripped, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, caption=True)
            continue
        # 关键词行
        if stripped.startswith('**关键词'):
            add_body(doc, stripped, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
            continue
        add_body(doc, stripped, indent=True)

    flush_bq(); flush_list(); flush_table()

    # 若至今未插入目录(异常情况)保证插入
    if toc_pending and not toc_emitted:
        doc.add_page_break()
        add_toc(doc)

    doc.save(OUT)
    print('SAVED:', OUT)

if __name__ == '__main__':
    main()
